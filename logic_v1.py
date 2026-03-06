import re
import fitz 
from docx import Document
import spacy 
import argparse
import os
import io
import pandas as pd
from PIL import Image, ImageDraw
import pytesseract 
from pytesseract import Output

# Set tesseract path
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Load NLP
nlp = spacy.load("en_core_web_sm")

class PIIRedactor:
    def __init__(self):
        self.patterns = {
            "EMAIL": r'[a-zA-Z0-9_.+-]+\s*@\s*[a-zA-Z0-9-]+\s*\.\s*[a-zA-Z0-9-.]+',
            "PHONE": r'\b(?:\+91|91)?[6-9]\d{9}\b',
            "PAN": r'[A-Z]{5}[0-9]{4}[A-Z]{1}',
            "AADHAAR": r'\b\d{4}\s?\d{4}\s?\d{4}\b',
            "IP_ADDRESS": r'\b(?:\d{1,3}\.){3}\d{1,3}\b'
        }

    def redact_text_string(self, text, mode="MASK"):
        """Helper to redact PII within a string while keeping its length similar."""
        if not text: return text
        
        # 1. Regex
        for label, pattern in self.patterns.items():
            if label == "EMAIL":
                emails = re.findall(pattern, text)
                for email in emails:
                    clean_email = email.replace(" ", "")
                    user, domain = clean_email.split('@')
                    masked = f"{user[0]}***@{domain}"
                    text = text.replace(email, masked)
            else:
                replacement = "***" if mode == "MASK" else "[REDACTED]"
                text = re.sub(pattern, replacement, text)
        
        # 2. NLP
        doc = nlp(text)
        for ent in reversed(doc.ents):
            if ent.label_ in ["PERSON", "GPE"]:
                replacement = "***" if mode == "MASK" else "[REDACTED]"
                text = text[:ent.start_char] + replacement + text[ent.end_char:]
        return text

    def redact_image_bytes(self, image_bytes):
        """Standard image OCR redaction."""
        img = Image.open(io.BytesIO(image_bytes)).convert("RGB")
        draw = ImageDraw.Draw(img)
        d = pytesseract.image_to_data(img, output_type=Output.DICT)
        for i in range(len(d['text'])):
            word = d['text'][i]
            # Simple check for OCR snippets
            is_pii = False
            for pattern in self.patterns.values():
                if re.search(pattern, word): is_pii = True
            
            if is_pii:
                x, y, w, h = d['left'][i], d['top'][i], d['width'][i], d['height'][i]
                draw.rectangle([x-2, y-2, x+w+2, y+h+2], fill="black")
        
        img_byte_arr = io.BytesIO()
        img.save(img_byte_arr, format='PNG')
        return img_byte_arr.getvalue()

    def process_docx_unified(self, docx_path, mode="MASK"):
        """Redacts BOTH text and images inside the same .docx file."""
        doc = Document(docx_path)
        
        # 1. Redact Text in Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                para.text = self.redact_text_string(para.text, mode)
        
        # 2. Redact Text in Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.text = self.redact_text_string(paragraph.text, mode)

        # 3. Redact Images
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                img_part = rel.target_part
                img_part._blob = self.redact_image_bytes(img_part.blob)
        
        output_path = docx_path.replace(".docx", "_redacted_final.docx")
        doc.save(output_path)
        return output_path

    def process_pdf_unified(self, pdf_path, mode="MASK"):
        """Redacts BOTH text and images inside the same .pdf file."""
        doc = fitz.open(pdf_path)
        for page in doc:
            # 1. Redact Text layer
            for label, pattern in self.patterns.items():
                text_instances = page.search_for(pattern)
                for inst in text_instances:
                    page.add_redaction_annot(inst, fill=(0, 0, 0))
            page.apply_redactions()

            # 2. Redact Image layer
            for img in page.get_images():
                xref = img[0]
                base_image = doc.extract_image(xref)
                redacted_bytes = self.redact_image_bytes(base_image["image"])
                doc.update_stream(xref, redacted_bytes)
        
        output_path = pdf_path.replace(".pdf", "_redacted_final.pdf")
        doc.save(output_path)
        doc.close()
        return output_path

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--file", required=True)
    parser.add_argument("--mode", default="MASK")
    args = parser.parse_args()

    redactor = PIIRedactor()
    file_path = args.file.strip()
    ext = os.path.splitext(file_path)[1].lower()

    print(f"[*] Processing: {file_path}")

    try:
        if ext == '.docx':
            final_file = redactor.process_docx_unified(file_path, args.mode)
            print(f" Success! Combined file saved to: {final_file}")
        elif ext == '.pdf':
            final_file = redactor.process_pdf_unified(file_path, args.mode)
            print(f" Success! Combined file saved to: {final_file}")
        elif ext in ['.png', '.jpg', '.jpeg']:
            out = file_path.replace(ext, f"_redacted{ext}")
            final_file = redactor.redact_image_file(file_path, out)
            print(f" Success! Redacted image saved to: {final_file}")
        else:
            print("Unsupported format for unified redaction.")
    except Exception as e:
        print(f" [!] Error: {e}")

if __name__ == "__main__":
    main()
