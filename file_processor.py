"""
file_processor.py — Extract text from multiple formats, run PII scan, rebuild sanitized file
- Images: black box redaction drawn over PII regions using Pillow + pytesseract OCR
- PDF/DOCX/CSV/TXT/JSON/SQL: text masking with XXX / [REDACTED]
"""
import io
import csv
import json
import re

import pdfplumber
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from PIL import Image, ImageDraw

from pii_engine import full_scan, build_pii_summary, REGEX_PATTERNS


def process_file(file_bytes: bytes, filename: str) -> tuple:
    ext = filename.rsplit(".", 1)[-1].lower()

    if ext == "pdf":
        return _process_pdf(file_bytes)
    elif ext == "docx":
        return _process_docx(file_bytes)
    elif ext in ("sql", "txt"):
        return _process_text(file_bytes)
    elif ext == "csv":
        return _process_csv(file_bytes)
    elif ext == "json":
        return _process_json(file_bytes)
    elif ext in ("png", "jpg", "jpeg"):
        return _process_image(file_bytes, ext)
    else:
        return _process_text(file_bytes)


# ── IMAGE (Black Box Redaction) ───────────────────────────────────

def _process_image(file_bytes: bytes, ext: str = "png") -> tuple:
    """
    Black-box redaction on images using OCR:
    1. Get word + line level bounding boxes via tesseract
    2. Reconstruct lines and scan each line for PII using regex
    3. Draw black boxes over every word in a line that contains PII
    4. Return redacted image
    """
    try:
        import pytesseract
        from pytesseract import Output

        image = Image.open(io.BytesIO(file_bytes)).convert("RGB")
        draw = ImageDraw.Draw(image)
        data = pytesseract.image_to_data(image, output_type=Output.DICT)

        all_detections = []
        n = len(data["text"])

        # ── Group words into lines by (block_num, par_num, line_num) ──
        lines = {}
        for i in range(n):
            word = data["text"][i].strip()
            if not word:
                continue
            key = (data["block_num"][i], data["par_num"][i], data["line_num"][i])
            if key not in lines:
                lines[key] = []
            lines[key].append(i)

        # ── IMAGE: Only black-box phone, card, email, cvv ──────────
        IMAGE_PII_TYPES = {"credit_card", "cvv", "phone", "us_phone", "email"}

        for key, indices in lines.items():
            line_text = " ".join(data["text"][i].strip() for i in indices)

            for pii_type, (pattern, masker) in REGEX_PATTERNS.items():
                if pii_type not in IMAGE_PII_TYPES:
                    continue

                matches = list(re.finditer(pattern, line_text, re.IGNORECASE))
                if not matches:
                    continue

                for match in matches:
                    matched_str = match.group()
                    try:
                        masked_val = masker(matched_str)
                    except Exception:
                        masked_val = "****"

                    # Find exact character positions of match in line_text
                    # and only box words that overlap with those positions
                    match_start = match.start()
                    match_end = match.end()

                    # Rebuild word positions in line_text
                    pos = 0
                    word_positions = []
                    for i in indices:
                        word = data["text"][i].strip()
                        if not word:
                            word_positions.append((i, -1, -1))
                            continue
                        word_start = line_text.find(word, pos)
                        word_end = word_start + len(word) if word_start >= 0 else -1
                        word_positions.append((i, word_start, word_end))
                        if word_start >= 0:
                            pos = word_start + len(word)

                    for word_idx, (i, wstart, wend) in enumerate(word_positions):
                        if wstart < 0 or wend < 0:
                            continue
                        # Only box if this word overlaps with the regex match
                        if wend <= match_start or wstart >= match_end:
                            continue

                        word = data["text"][i].strip()

                        # credit_card: keep first group (first 4 digits) visible
                        if pii_type == "credit_card" and wstart == match_start:
                            continue

                        # phone: keep country code visible
                        if pii_type in ("phone", "us_phone"):
                            if word.startswith("+") or word in ("+91", "+1"):
                                continue

                        # email: keep @domain part visible
                        if pii_type == "email":
                            at_pos = matched_str.find("@")
                            if at_pos >= 0 and wstart >= match_start + at_pos:
                                continue

                        x = data["left"][i]
                        y = data["top"][i]
                        w = data["width"][i]
                        h = data["height"][i]
                        draw.rectangle([x-3, y-3, x+w+3, y+h+3], fill="black")



                    all_detections.append({
                        "pii_type": pii_type,
                        "original_value": matched_str,
                        "masked_value": masked_val,
                        "detection_method": "regex+ocr",
                        "confidence": 1.0
                    })

        # Save redacted image
        out = io.BytesIO()
        fmt = "JPEG" if ext in ("jpg", "jpeg") else "PNG"
        image.save(out, format=fmt)
        summary = build_pii_summary(all_detections)
        return out.getvalue(), all_detections, summary

    except ImportError:
        # pytesseract not installed
        image = Image.open(io.BytesIO(file_bytes)).convert("RGB")
        draw = ImageDraw.Draw(image)
        iw, ih = image.size
        draw.rectangle([0, ih//2 - 40, iw, ih//2 + 40], fill=(150, 0, 0))
        draw.text((20, ih//2 - 15), "Install tesseract for OCR image redaction", fill="white")
        out = io.BytesIO()
        image.save(out, format="PNG")
        return out.getvalue(), [], {"error": "pytesseract not installed"}

    except Exception as e:
        print(f"[Image processing error] {e}")
        return file_bytes, [], {}


# ── PDF ──────────────────────────────────────────────────────────

def _process_pdf(file_bytes: bytes) -> tuple:
    all_detections = []
    full_masked_text = ""

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            masked, detections = full_scan(text)
            all_detections.extend(detections)
            full_masked_text += masked + "\n\n"

    # Rebuild as sanitized PDF
    out = io.BytesIO()
    doc = SimpleDocTemplate(out, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []
    for line in full_masked_text.split("\n"):
        if line.strip():
            try:
                safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(safe, styles["Normal"]))
            except Exception:
                pass
    doc.build(story)
    summary = build_pii_summary(all_detections)
    return out.getvalue(), all_detections, summary


# ── DOCX ─────────────────────────────────────────────────────────

def _process_docx(file_bytes: bytes) -> tuple:
    doc = Document(io.BytesIO(file_bytes))
    all_detections = []

    for para in doc.paragraphs:
        for run in para.runs:
            if run.text.strip():
                masked, detections = full_scan(run.text)
                all_detections.extend(detections)
                run.text = masked

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text.strip():
                            masked, detections = full_scan(run.text)
                            all_detections.extend(detections)
                            run.text = masked

    out = io.BytesIO()
    doc.save(out)
    summary = build_pii_summary(all_detections)
    return out.getvalue(), all_detections, summary


# ── SQL / TXT ────────────────────────────────────────────────────

def _process_text(file_bytes: bytes) -> tuple:
    text = file_bytes.decode("utf-8", errors="replace")
    masked, detections = full_scan(text)
    summary = build_pii_summary(detections)
    return masked.encode("utf-8"), detections, summary


# ── CSV ──────────────────────────────────────────────────────────

def _process_csv(file_bytes: bytes) -> tuple:
    text = file_bytes.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    rows = list(reader)
    all_detections = []

    sanitized_rows = []
    for row in rows:
        sanitized_row = []
        for cell in row:
            if cell.strip():
                masked, detections = full_scan(cell)
                all_detections.extend(detections)
                sanitized_row.append(masked)
            else:
                sanitized_row.append(cell)
        sanitized_rows.append(sanitized_row)

    out = io.StringIO()
    writer = csv.writer(out)
    writer.writerows(sanitized_rows)
    summary = build_pii_summary(all_detections)
    return out.getvalue().encode("utf-8"), all_detections, summary


# ── JSON ─────────────────────────────────────────────────────────

def _process_json(file_bytes: bytes) -> tuple:
    text = file_bytes.decode("utf-8", errors="replace")
    masked, detections = full_scan(text)
    summary = build_pii_summary(detections)
    return masked.encode("utf-8"), detections, summary


# ── PREVIEW ──────────────────────────────────────────────────────

def extract_preview_text(file_bytes: bytes, filename: str, max_chars: int = 2000) -> str:
    ext = filename.rsplit(".", 1)[-1].lower()
    try:
        if ext == "pdf":
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)
                return text[:max_chars]
        elif ext == "docx":
            doc = Document(io.BytesIO(file_bytes))
            return "\n".join(p.text for p in doc.paragraphs)[:max_chars]
        else:
            return file_bytes.decode("utf-8", errors="replace")[:max_chars]
    except Exception:
        return "[Preview unavailable]"
